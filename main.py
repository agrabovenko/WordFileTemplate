import tkinter as tk
from tkinter import messagebox, filedialog
from docx import Document
import os
import random


template_path = None


def select_template():
    global template_path

    file_path = filedialog.askopenfilename(
        title="Выберите шаблон документа",
        filetypes=[("Word files", "*.docx")]
    )

    if not file_path:
        return

    template_path = file_path
    template_label.config(text=f"Шаблон: {os.path.basename(file_path)}")


def generate_doc():
    if not template_path:
        messagebox.showwarning("Ошибка", "Сначала выбери шаблон!")
        return

    name = name_input.get()
    age = age_input.get()
    date = date_input.get()
    filename = file_input.get()
    gander = gender_input.get()
    some_key = "".join(random.choice("01") for _ in range(20))

    if not name or not age or not date or not filename:
        messagebox.showwarning("Ошибка", "Заполни все поля!")
        return

    values = {
        "{{NAME}}": name,
        "{{AGE}}": age,
        "{{GENDER}}": gander,
        "{{DATE}}": date,
        "{{KEY}}": some_key
    }

    doc = Document(template_path)

    for p in doc.paragraphs:
        for key, val in values.items():
            if key in p.text:
                p.text = p.text.replace(key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in values.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

    output_name = f"{filename}.docx"
    doc.save(output_name)

    messagebox.showinfo("Готово", f"Документ сохранён как:\n{output_name}")


root = tk.Tk()
root.title("Генератор документов")
root.geometry("680x420")


tk.Button(root, text="Выбрать шаблон", command=select_template).pack(pady=5)
template_label = tk.Label(root, text="Шаблон не выбран")
template_label.pack()


tk.Label(root, text="Дата:").pack(anchor="w", padx=10)
date_input = tk.Entry(root)
date_input.pack(fill="x", padx=10)


tk.Label(root, text="Имя:").pack(anchor="w", padx=10)
name_input = tk.Entry(root)
name_input.pack(fill="x", padx=10)


tk.Label(root, text="Пол:").pack(anchor="w", padx=10)
gender_input = tk.Entry(root)
gender_input.pack(fill="x", padx=10)


tk.Label(root, text="Возраст:").pack(anchor="w", padx=10)
age_input = tk.Entry(root)
age_input.pack(fill="x", padx=10)


tk.Label(root, text="Название файла (без .docx):").pack(anchor="w", padx=10)
file_input = tk.Entry(root)
file_input.pack(fill="x", padx=10)


tk.Button(root, text="Создать документ", command=generate_doc).pack(pady=15)

root.mainloop()